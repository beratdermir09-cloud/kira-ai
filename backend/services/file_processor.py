import io
import json
import csv as csv_module
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF — pdfplumber önce, PyPDF2 fallback."""
    # pdfplumber dene (daha iyi metin çıkarımı)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"--- Sayfa {i+1} ---\n{text.strip()}")
            if pages:
                return "\n\n".join(pages)
    except ImportError:
        pass
    except Exception:
        pass

    # PyPDF2 fallback
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Sayfa {i+1} ---\n{text.strip()}")
        if pages:
            return "\n\n".join(pages)
        return "⚠️ PDF'den metin çıkarılamadı. Dosya taranmış görsel içeriyor olabilir."
    except Exception as e:
        return f"PDF okuma hatası: {str(e)}"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX — paragraflar + tablolar."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # Paragraflar
        for para in doc.paragraphs:
            if para.text.strip():
                # Başlık stillerini işaretle
                if para.style.name.startswith('Heading'):
                    parts.append(f"\n## {para.text.strip()}")
                else:
                    parts.append(para.text.strip())

        # Tablolar
        for i, table in enumerate(doc.tables):
            parts.append(f"\n[Tablo {i+1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip() or "Döküman boş görünüyor."
    except Exception as e:
        return f"DOCX okuma hatası: {str(e)}"


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text files."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1254"):
        try:
            return file_bytes.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


def extract_text_from_csv(file_bytes: bytes) -> str:
    """Extract and format CSV content."""
    try:
        text = extract_text_from_txt(file_bytes)
        reader = csv_module.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return "CSV dosyası boş."

        # İlk satır başlık
        header = rows[0]
        total_rows = len(rows) - 1
        preview_rows = rows[1:51]  # İlk 50 satır

        lines = [f"Toplam {total_rows} satır, {len(header)} sütun\n"]
        lines.append(" | ".join(header))
        lines.append("-" * min(80, len(" | ".join(header))))
        for row in preview_rows:
            lines.append(" | ".join(row))
        if total_rows > 50:
            lines.append(f"\n... ve {total_rows - 50} satır daha")
        return "\n".join(lines)
    except Exception:
        return extract_text_from_txt(file_bytes)


def extract_text_from_json(file_bytes: bytes) -> str:
    """Extract and pretty-print JSON content."""
    try:
        text = extract_text_from_txt(file_bytes)
        data = json.loads(text)
        pretty = json.dumps(data, ensure_ascii=False, indent=2)
        # Çok büyükse kısalt
        if len(pretty) > 15000:
            pretty = pretty[:15000] + "\n... (JSON çok büyük, kısaltıldı)"
        return pretty
    except Exception:
        return extract_text_from_txt(file_bytes)


def extract_text_from_excel(file_bytes: bytes) -> str:
    """Extract text from Excel file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"[Sayfa: {sheet_name}]\n" + "\n".join(rows[:100]))
                if len(rows) > 100:
                    parts.append(f"... ve {len(rows)-100} satır daha")
        return "\n\n".join(parts) or "Excel dosyası boş."
    except ImportError:
        pass
    except Exception as e:
        return f"Excel okuma hatası: {str(e)}"

    # pandas fallback
    try:
        import pandas as pd
        xl = pd.ExcelFile(io.BytesIO(file_bytes))
        parts = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            parts.append(f"[Sayfa: {sheet}]\n{df.to_string(index=False, max_rows=100)}")
        return "\n\n".join(parts) or "Excel dosyası boş."
    except Exception as e:
        return f"Excel okuma hatası: {str(e)}"


def process_file(filename: str, file_bytes: bytes) -> str:
    """Process uploaded file and return extracted text."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return extract_text_from_excel(file_bytes)
    elif ext == "csv":
        return extract_text_from_csv(file_bytes)
    elif ext == "json":
        return extract_text_from_json(file_bytes)
    elif ext in ("txt", "md", "py", "js", "ts", "jsx", "tsx", "html", "css",
                 "xml", "yaml", "yml", "toml", "ini", "env", "sh", "bat",
                 "sql", "rs", "go", "java", "c", "cpp", "h", "cs", "php", "rb"):
        return extract_text_from_txt(file_bytes)
    else:
        # Bilinmeyen format — yine de metin olarak okumayı dene
        try:
            return extract_text_from_txt(file_bytes)
        except Exception:
            return f"Desteklenmeyen dosya formatı: .{ext}"


def truncate_content(content: str, max_chars: int = 20000) -> str:
    """Truncate content to avoid token limits — akıllı kırpma."""
    if len(content) <= max_chars:
        return content

    # Sayfa sınırlarında kesmek için --- Sayfa N --- işaretlerini kullan
    half = max_chars // 2
    head = content[:half]
    tail = content[-half:]

    # Satır ortasında kesme
    head = head[:head.rfind('\n') + 1] if '\n' in head else head
    tail = tail[tail.find('\n') + 1:] if '\n' in tail else tail

    skipped = len(content) - len(head) - len(tail)
    return (
        head
        + f"\n\n[... {skipped:,} karakter atlandı — dosya çok uzun ...]\n\n"
        + tail
    )
